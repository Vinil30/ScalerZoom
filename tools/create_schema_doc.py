from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_DOCX = DOCS / "AI_Zoom_SQL_Schema_Diagram.docx"
OUT_PNG = DOCS / "schema_diagram.png"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


FONT_TITLE = font(34, True)
FONT_SUBTITLE = font(20)
FONT_BOX_TITLE = font(22, True)
FONT_FIELD = font(17)
FONT_SMALL = font(14)
FONT_BADGE = font(15, True)


def rounded_box(draw: ImageDraw.ImageDraw, box, fill, outline, radius=18, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_text_lines(draw, xy, text, fnt, fill, max_chars, line_gap=5):
    x, y = xy
    for line in wrap(text, max_chars):
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def draw_table_box(draw, x, y, w, title, fields, fill, outline, accent):
    row_h = 28
    header_h = 44
    h = header_h + row_h * len(fields) + 18
    rounded_box(draw, (x, y, x + w, y + h), fill, outline, radius=18, width=3)
    draw.rounded_rectangle((x, y, x + w, y + header_h), radius=18, fill=accent)
    draw.rectangle((x, y + header_h - 14, x + w, y + header_h), fill=accent)
    draw.text((x + 18, y + 10), title, font=FONT_BOX_TITLE, fill="white")
    yy = y + header_h + 10
    for field in fields:
        draw.text((x + 18, yy), field, font=FONT_FIELD, fill="#273142")
        yy += row_h
    return (x, y, x + w, y + h)


def arrow(draw, start, end, color="#64748b", width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - direction * 18, ey - 8), (ex - direction * 18, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 8, ey - direction * 18), (ex + 8, ey - direction * 18)]
    draw.polygon(points, fill=color)


def badge(draw, x, y, text, fill="#eff6ff", outline="#bfdbfe", color="#1d4ed8"):
    tw = draw.textlength(text, font=FONT_BADGE)
    rounded_box(draw, (x, y, x + tw + 28, y + 34), fill, outline, radius=14, width=2)
    draw.text((x + 14, y + 8), text, font=FONT_BADGE, fill=color)


def create_diagram():
    img = Image.new("RGB", (2200, 1450), "#f8fafc")
    draw = ImageDraw.Draw(img)

    draw.text((70, 50), "AI Zoom Clone - SQL Database Schema", font=FONT_TITLE, fill="#0f172a")
    draw.text(
        (70, 96),
        "Sketch-style relationship map with tables, feature flows, and API touchpoints",
        font=FONT_SUBTITLE,
        fill="#475569",
    )

    users = draw_table_box(
        draw,
        80,
        190,
        430,
        "users",
        ["PK id", "username UNIQUE", "email UNIQUE", "avatar_url", "created_at / updated_at"],
        "#ffffff",
        "#cbd5e1",
        "#2563eb",
    )
    meetings = draw_table_box(
        draw,
        820,
        175,
        500,
        "meetings",
        ["PK id", "meeting_uuid UNIQUE", "meeting_code UNIQUE", "FK host_id -> users.id", "title / description", "meeting_type CHECK", "status CHECK", "scheduled_start", "duration_minutes"],
        "#ffffff",
        "#cbd5e1",
        "#0f766e",
    )
    participants = draw_table_box(
        draw,
        1580,
        190,
        500,
        "participants",
        ["PK id", "FK meeting_id -> meetings.id", "FK user_id -> users.id NULL", "display_name", "role CHECK", "joined_at / left_at", "mic_enabled / video_enabled", "UNIQUE(meeting_id, user_id)"],
        "#ffffff",
        "#cbd5e1",
        "#7c3aed",
    )
    links = draw_table_box(
        draw,
        110,
        720,
        440,
        "meeting_links",
        ["PK id", "FK meeting_id -> meetings.id", "invite_link UNIQUE", "created_at", "expires_at"],
        "#ffffff",
        "#cbd5e1",
        "#0284c7",
    )
    history = draw_table_box(
        draw,
        650,
        760,
        420,
        "meeting_history",
        ["PK id", "FK meeting_id -> meetings.id", "participant_count", "started_at / ended_at", "total_duration"],
        "#ffffff",
        "#cbd5e1",
        "#ea580c",
    )
    transcripts = draw_table_box(
        draw,
        1190,
        710,
        430,
        "ai_transcripts",
        ["PK id", "FK meeting_id -> meetings.id", "transcript_text", "language", "source_model", "created_at"],
        "#ffffff",
        "#cbd5e1",
        "#0891b2",
    )
    summaries = draw_table_box(
        draw,
        1660,
        690,
        430,
        "ai_meeting_summaries",
        ["PK id", "FK meeting_id -> meetings.id", "generated_summary", "generated_by_model", "created_at"],
        "#ffffff",
        "#cbd5e1",
        "#16a34a",
    )
    actions = draw_table_box(
        draw,
        1360,
        990,
        500,
        "ai_action_items",
        ["PK id", "FK meeting_id -> meetings.id", "action_text", "assigned_to", "priority CHECK", "status CHECK", "generated_at"],
        "#ffffff",
        "#cbd5e1",
        "#dc2626",
    )

    # Relationship links
    arrow(draw, (users[2], 315), (meetings[0], 315), "#2563eb")
    draw.text((555, 286), "hosts", font=FONT_SMALL, fill="#1d4ed8")
    arrow(draw, (users[2], 420), (participants[0], 420), "#7c3aed")
    draw.text((600, 432), "user joins", font=FONT_SMALL, fill="#6d28d9")
    arrow(draw, (meetings[2], 360), (participants[0], 360), "#64748b")
    draw.text((1350, 332), "meeting has participants", font=FONT_SMALL, fill="#475569")
    arrow(draw, (930, meetings[3]), (links[2], links[1]), "#0284c7")
    arrow(draw, (990, meetings[3]), (history[0] + 100, history[1]), "#ea580c")
    arrow(draw, (1060, meetings[3]), (transcripts[0] + 90, transcripts[1]), "#0891b2")
    arrow(draw, (1190, meetings[3]), (summaries[0] + 80, summaries[1]), "#16a34a")
    arrow(draw, (1270, meetings[3]), (actions[0] + 120, actions[1]), "#dc2626")

    # Feature layer
    rounded_box(draw, (80, 1300, 2080, 1405), "#ffffff", "#dbeafe", radius=22, width=3)
    draw.text((110, 1320), "Feature Links", font=FONT_BOX_TITLE, fill="#0f172a")
    badge(draw, 110, 1362, "Dashboard -> meetings + participants + AI counts")
    badge(draw, 590, 1362, "Join Flow -> meeting_code + participants")
    badge(draw, 980, 1362, "AI Pipeline -> transcripts -> summaries -> action items")
    badge(draw, 1500, 1362, "Meeting Room -> media state + participant sync")

    # Notes
    rounded_box(draw, (1420, 170, 2075, 535), "#f8fafc", "#bfdbfe", radius=18, width=2)
    draw.text((1450, 198), "API Touchpoints", font=FONT_BOX_TITLE, fill="#1e3a8a")
    api_lines = [
        "POST /api/v1/meetings",
        "POST /api/v1/meetings/join",
        "POST /api/v1/schedule",
        "GET /api/v1/dashboard/overview",
        "POST /api/v1/ai/transcripts/process",
        "GET /api/v1/ai/meetings/{id}/action-items",
    ]
    yy = 242
    for line in api_lines:
        draw.text((1450, yy), line, font=FONT_FIELD, fill="#334155")
        yy += 34

    img.save(OUT_PNG, quality=95)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_docx():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(9.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Zoom Clone - SQL Schema Diagram")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run("Visual database sketch with tables, relationships, feature links, and API touchpoints")
    sub.font.size = Pt(10.5)
    sub.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_picture(str(OUT_PNG), width=Inches(10.75))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Reading guide: meetings is the central table; users host/join; AI artifacts attach to meetings as durable generated knowledge.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_page_break()

    h = doc.add_paragraph()
    hr = h.add_run("Feature-to-Table Map")
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.color.rgb = RGBColor(15, 23, 42)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Feature", "Primary Tables", "Backend API Links", "Why It Matters"]
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "DBEAFE")
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="1E3A8A")

    rows = [
        ("Dashboard", "meetings, participants, ai_transcripts, ai_meeting_summaries", "GET /api/v1/dashboard/overview", "Shows meeting health, participant activity, and AI usage at a glance."),
        ("Instant Meeting", "meetings, meeting_links, participants, meeting_history", "POST /api/v1/meetings", "Creates a live room, invite code, host participant, and history row."),
        ("Join Meeting", "meetings, meeting_links, participants", "POST /api/v1/meetings/join", "Validates codes, expired links, duplicate joins, and participant state."),
        ("Scheduling", "meetings, meeting_links", "POST /api/v1/schedule", "Persists future meetings and feeds upcoming dashboard sections."),
        ("AI Summary", "ai_transcripts, ai_meeting_summaries", "POST /api/v1/ai/summaries/generate", "Keeps generated recaps auditable and linked to source meetings."),
        ("AI Action Items", "ai_transcripts, ai_action_items", "POST /api/v1/ai/action-items/generate", "Turns meeting discussion into persisted tasks with owner and priority."),
        ("Transcript Intelligence", "ai_transcripts, ai_meeting_summaries, ai_action_items", "POST /api/v1/ai/transcripts/process", "One workflow stores transcript, summary, and extracted tasks."),
        ("Media State", "participants", "PATCH /api/v1/participants/{id}", "Synchronizes mic/camera state without overbuilding streaming infrastructure."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    r2 = h2.add_run("Interview Talking Points")
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(15, 23, 42)

    bullets = [
        "The schema is normalized around real product concepts, not CRUD screens.",
        "Meeting-owned records cascade so cleanup is predictable.",
        "AI tables are separate because generated artifacts can be regenerated and compared by model.",
        "Indexes match actual API access patterns: join code lookup, dashboard filtering, and latest AI artifact retrieval.",
        "SQLite keeps the assignment deployable while preserving a design that can migrate to PostgreSQL.",
    ]
    for item in bullets:
        para = doc.add_paragraph(style=None)
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.15)
        para.add_run("• ").bold = True
        para.add_run(item)

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    DOCS.mkdir(exist_ok=True)
    create_diagram()
    create_docx()
    print(OUT_DOCX)
